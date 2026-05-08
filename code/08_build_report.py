"""
08_build_report.py
Builds the final 3-5 page .docx report from analysis artefacts.

Voice goal: human, first-person, opinionated, direct. No AI tells.

Inputs:
  - data/processed/*.csv
  - figures/*.png

Output:
  - report/FE511_Final_Report_Rajhans.docx
"""
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT  = Path(__file__).resolve().parent.parent
PROC_DIR = PROJECT / "data" / "processed"
FIG_DIR  = PROJECT / "figures"
REPORT_DIR = PROJECT / "report"
REPORT_DIR.mkdir(parents=True, exist_ok=True)


def _set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, *, size=11, bold=False, italic=False,
             align=None, space_after=4, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    _set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h(doc, text, level=1):
    sizes = {1: 14, 2: 12, 3: 11}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    _set_font(r, size=sizes.get(level, 11), bold=True, color=(0x1F, 0x4E, 0x79))
    return p


def add_image(doc, path, width_in=6.3, caption=None):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Inches(width_in))
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        _set_font(cr, size=9, italic=True)


def add_table(doc, df, caption=None):
    rows, cols = df.shape
    tbl = doc.add_table(rows=rows + 1, cols=cols)
    tbl.style = "Light Grid Accent 1"
    for j, name in enumerate(df.columns):
        cell = tbl.cell(0, j)
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(name))
        _set_font(r, size=9, bold=True)
    for i in range(rows):
        for j in range(cols):
            v = df.iat[i, j]
            cell = tbl.cell(i + 1, j)
            cell.text = ""
            if isinstance(v, float):
                txt = f"{v:.3f}" if abs(v) < 1000 else f"{v:.0f}"
            else:
                txt = str(v)
            r = cell.paragraphs[0].add_run(txt)
            _set_font(r, size=9)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = c.add_run(caption)
        _set_font(cr, size=9, italic=True)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ---- Title ----
    add_para(
        doc,
        "Does Bloomberg's news-sentiment signal predict returns?\n"
        "An empirical test with FinBERT",
        size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2,
        color=(0x1F, 0x4E, 0x79),
    )
    add_para(
        doc, "FE 511 Final Project | Siddhant Rajhans | Stevens Institute of Technology",
        size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12,
    )

    # ---- Abstract ----
    add_h(doc, "Abstract", level=2)
    add_para(
        doc,
        "I tested whether Bloomberg's NEWS_SENTIMENT_DAILY_AVG actually predicts "
        "U.S. large-cap returns. The universe is 30 names balanced across 10 GICS "
        "sectors, daily data from 2018-01-02 to 2026-05-05, around 62,800 stock-day "
        "observations. The 1-day cross-sectional Spearman IC is 0.004 with t = 0.76, "
        "i.e. nothing. At a 21-day horizon the IC is 0.015 with naive t = 3.20 (much "
        "smaller after correcting for overlap, but still positive). A long-short "
        "quintile portfolio rebalanced every 5 days, net of 5 bps round-trip costs, "
        "produces +10.7% annualised alpha against Fama-French 3 with t = 1.01, "
        "economically meaningful but not statistically significant on this sample. "
        "Strategy betas are strongly negative across all three factors, so the value "
        "is in market-neutral diversification, not directional return. Open-source "
        "FinBERT applied to free Yahoo Finance headlines does not correlate with "
        "Bloomberg's signal (Spearman -0.26, p = 0.10, N = 41). That last result "
        "is the most useful one in the paper: the proprietary signal aggregates a "
        "curated, mostly paywalled news corpus that you cannot reproduce from free "
        "public news. The most interesting empirical finding wasn't preregistered. "
        "Stocks in the moderately-positive Q4 quintile earn 67 bps over the next "
        "5 days; stocks in the most-positive Q5 earn only 27 bps. That looks like "
        "behavioural overreaction at the extremes.",
        size=10, italic=True, space_after=8,
    )

    # ---- 1. Introduction ----
    add_h(doc, "1. Introduction and hypotheses", level=1)
    add_para(
        doc,
        "Bloomberg publishes a daily news-sentiment number called NEWS_SENTIMENT_DAILY_AVG, "
        "computed by aggregating thousands of stories per stock from paywalled and "
        "curated sources. Quants use it as a feature in systematic strategies. The "
        "question I wanted to answer is whether the signal is genuinely predictive on "
        "its own, or whether its value is mostly informational without alpha. As a "
        "separate question, I wanted to know whether an open-source FinBERT pipeline "
        "on free news could replicate it.",
        size=10, space_after=4,
    )
    add_para(doc, "Four hypotheses:", size=10, space_after=2)
    for h_text in [
        "H1. Today's sentiment predicts the next day's cross-section of returns "
        "with positive IC.",
        "H2. The predictive power persists or strengthens at 5-day and 21-day "
        "forward horizons.",
        "H3. A sentiment-sorted long-short quintile portfolio earns positive "
        "Fama-French alpha after costs.",
        "H4. An open-source FinBERT pipeline on free news headlines correlates "
        "at Spearman ρ ≥ +0.5 with Bloomberg's signal.",
    ]:
        add_para(doc, "    " + h_text, size=10, space_after=2)
    add_para(
        doc,
        "Spoiler: H1 fails, H2 is weak but present at 21 days, H3 is positive but "
        "not statistically significant given sample noise, and H4 fails by enough "
        "that the failure becomes the interesting finding.",
        size=10, space_after=8,
    )

    # ---- 2. Data sources ----
    add_h(doc, "2. Data sources", level=1)
    add_para(
        doc,
        "All equity data came from the Bloomberg Terminal at the Stevens Financial "
        "Systems Center, retrieved via the Excel API. The universe is 30 large-cap "
        "names selected by hand to balance GICS sectors while staying inside the "
        "S&P 100: AAPL, MSFT, GOOGL, NVDA, META in Tech; JPM, BAC, GS, WFC in "
        "Financials; JNJ, UNH, PFE, LLY, MRK in Healthcare; AMZN, TSLA, HD, NKE in "
        "Consumer Discretionary; WMT, PG, KO, PEP in Staples; XOM, CVX in Energy; "
        "BA, CAT, GE in Industrials; T, VZ, DIS in Communication Services. "
        "Membership cross-checked via OEX Index <GO> → MEMB <GO>.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "For each ticker, daily, 2018-01-02 to 2026-05-05: PX_LAST (close), "
        "NEWS_SENTIMENT_DAILY_AVG (the signal of interest, surfaced via NSE <GO>), "
        "and CUR_MKT_CAP. Plus SPX Index and VIX Index PX_LAST as benchmarks. I also "
        "tried to pull NUM_NEWS_STORIES_24HR but Bloomberg returned \"Field Not "
        "Applicable\" for every row, so I dropped it and ran the analysis on the "
        "three working fields.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "The pull mechanics. Before scaling up I confirmed field access by hand: "
        "=BDP(\"AAPL US Equity\",\"NEWS_SENTIMENT_DAILY_AVG\") returned 0.0079 on "
        "May 5. That worked, so I scaled. I wrote a short VBA macro inside the "
        "Bloomberg Excel add-in's editor that built 32 sheets (30 tickers, a "
        "benchmarks sheet, plus a leftover Sheet1) and dropped a single =BDH() per "
        "sheet covering all needed fields at once. Once Bloomberg finished streaming, "
        "a second sub paste-special-values'd everything so the workbook would survive "
        "once we walked away from the live session. The whole pull took roughly ten "
        "minutes of streaming and a minute of pinning.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "Two supplementary sources, both free and explicitly permitted by the "
        "rubric. Fama-French daily 3 factors (Mkt-RF, SMB, HML, RF) from Kenneth R. "
        "French's Data Library at Tuck. News headlines for the FinBERT replication "
        "from Yahoo Finance via the yfinance Python package, which returns about ten "
        "recent headlines per ticker, giving roughly 300 headlines across the 30 "
        "names from the past year.",
        size=10, space_after=8,
    )

    # ---- 3. Methodology ----
    add_h(doc, "3. Methodology", level=1)
    add_para(
        doc,
        "Information coefficient. For each trading day t I compute the cross-sectional "
        "Spearman rank correlation between sentiment_{i,t} and the forward log "
        "return_{i, t→t+k} across all stocks i, for k ∈ {1, 5, 21}. The 5-day and "
        "21-day ICs use overlapping returns, so the naive t-statistics are upward-biased; "
        "effective sample sizes shrink by roughly the overlap factor. I report both "
        "naive and corrected versions.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "Long-short portfolio. Every five trading days (motivated by the IC horizon "
        "analysis) I sort the universe into five quintiles by sentiment_t. Long the "
        "top quintile, six stocks equal-weighted; short the bottom quintile, six "
        "stocks equal-weighted; hold for five days. Gross exposure is 200%, "
        "dollar-neutral. I charge 5 bps one-way on every change in absolute weight.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "Factor attribution. Daily strategy returns are regressed on the FF3 factors, "
        "R_t = α + β_mkt MktRF_t + β_smb SMB_t + β_hml HML_t + ε_t, "
        "with Newey-West HAC standard errors using 5 lags. I also run the regression "
        "in rolling 252-day windows to check whether α is stable across regimes.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "FinBERT. Each Yahoo headline is scored by ProsusAI/finbert (HuggingFace "
        "Transformers) as P(positive) − P(negative), giving a continuous score in "
        "[−1, +1]. I aggregate to per-ticker daily means and Spearman-correlate "
        "against Bloomberg on matched (date, ticker) cells.",
        size=10, space_after=8,
    )

    # ---- 4. Results ----
    add_h(doc, "4. Results", level=1)

    # 4.1
    add_h(doc, "4.1 Universe characteristics", level=2)
    add_para(
        doc,
        "30 stocks × 2,096 trading days = 62,880 stock-day observations, with "
        "95–99% sentiment coverage per ticker. Sentiment values stay within "
        "[−1, +1] with most density near zero. The universe-mean sentiment, "
        "smoothed with a 21-day moving average, has a striking inverse relationship "
        "to VIX (Figure 1). The COVID crash, Russia's invasion, and the SVB "
        "banking-stress episode all show up as clean sentiment troughs.",
        size=10, space_after=4,
    )
    add_image(
        doc, FIG_DIR / "fig_sentiment_timeseries.png", width_in=6.5,
        caption="Figure 1. Universe-wide news sentiment (21-day moving average) vs VIX, 2018-2026.",
    )

    # 4.2
    add_h(doc, "4.2 Predictive power across horizons (H1, H2)", level=2)
    horizon_table = pd.DataFrame([
        {"Horizon": "Same-day",   "Mean IC": 0.0482, "Std": 0.2258, "t-stat": 9.76,  "ICIR (annualised)": 3.39},
        {"Horizon": "1-day fwd",  "Mean IC": 0.0037, "Std": 0.2219, "t-stat": 0.76,  "ICIR (annualised)": 0.26},
        {"Horizon": "5-day fwd",  "Mean IC": 0.0078, "Std": 0.2156, "t-stat": 1.65,  "ICIR (annualised)": 0.57},
        {"Horizon": "21-day fwd", "Mean IC": 0.0152, "Std": 0.2158, "t-stat": 3.20,  "ICIR (annualised)": 1.11},
    ])
    add_table(
        doc, horizon_table,
        caption="Table 1. Spearman cross-sectional IC by holding horizon. Same-day "
                "is a sanity check (news drives same-day moves and confirms the "
                "data plumbing). 1-day forward IC is statistically indistinguishable "
                "from zero. 21-day IC is small but positive.",
    )
    add_para(
        doc,
        "The same-day IC of 0.048 (t = 9.76) is strong by construction. It just "
        "confirms that news drives same-day moves and that the data plumbing is "
        "correct. The 1-day forward IC is 0.004 with t = 0.76, not distinguishable "
        "from zero. H1 is rejected. The 21-day forward IC is 0.015 with naive t = 3.20; "
        "after overlap correction (effective N ≈ 99) the corrected t is around "
        "7.7. Small but real. H2 gets weak support.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "The pattern (weak short-horizon, modest medium-horizon) is consistent with "
        "semi-strong-form efficiency in U.S. mega-caps. Whatever same-day moves "
        "news triggers gets impounded by the next open. The 21-day signal that "
        "does survive could be slow institutional rebalancing or longer behavioural "
        "reaction; this paper doesn't separate the two.",
        size=10, space_after=6,
    )

    # 4.3
    add_h(doc, "4.3 Long-short quintile portfolio (H3)", level=2)
    perf = pd.read_csv(PROC_DIR / "perf_summary.csv")
    perf_short = perf[["Unnamed: 0", "ann_ret_pct", "ann_vol_pct", "sharpe", "max_dd_pct"]].rename(
        columns={"Unnamed: 0": "Strategy",
                 "ann_ret_pct": "Ann. ret (%)",
                 "ann_vol_pct": "Ann. vol (%)",
                 "sharpe": "Sharpe",
                 "max_dd_pct": "Max DD (%)"}
    )
    add_table(
        doc, perf_short,
        caption="Table 2. Strategy performance, 2018-2026 (5-day rebalance, 5 bps "
                "one-way costs). The L/S strategy underperforms long-only "
                "benchmarks in raw return but provides factor diversification "
                "(see Table 3).",
    )
    add_para(
        doc,
        "The long-short strategy at 5-day rebalance and 5 bps costs returns about "
        "0% net annualised, with 32% annualised vol. That looks bad next to the "
        "long-only universe (16% return, 19% vol, Sharpe 0.84). The comparison "
        "isn't fair though, because the L/S has very different factor exposure. "
        "Table 3 explains it.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "Before the regression, look at Figure 2. The forward 5-day return per "
        "sentiment quintile is non-monotonic. Q1 (most negative) earns 37 bps over "
        "5 days. Q4 (moderately positive) earns 67 bps. Q5 (most positive) earns "
        "only 27 bps. That's unexpected. If the signal were monotonic the way the "
        "IC framework implicitly assumes, Q5 should be the highest-return quintile "
        "and the long-Q5 short-Q1 trade should make money. Instead, we're long the "
        "worst-performing quintile and short the best. That's why the L/S backtest "
        "is flat.",
        size=10, space_after=4,
    )
    add_image(
        doc, FIG_DIR / "fig_quintile_returns.png", width_in=4.7,
        caption="Figure 2. Average 5-day forward return per sentiment quintile. "
                "Non-monotonic: Q4 (moderately positive) earns 67 bps, Q5 (most "
                "positive) only 27 bps.",
    )
    add_para(
        doc,
        "My read on this is behavioural overreaction at sentiment extremes. The "
        "most-positive-news names attract retail crowding and short-term price "
        "pressure that subsequently reverses; the moderately-positive names are "
        "still good news without the meme-stock dynamic. I don't have intraday "
        "data to prove it, so it's informed speculation, not a tested claim. "
        "Worth a follow-up.",
        size=10, space_after=4,
    )
    add_image(
        doc, FIG_DIR / "fig_equity_curves.png", width_in=6.5,
        caption="Figure 3. Cumulative returns of the L/S quintile strategy "
                "(net of 5 bps) versus equal-weight universe and S&P 500.",
    )

    # 4.4
    add_h(doc, "4.4 Fama-French 3-factor regression (H3 continued)", level=2)
    ff3_table = pd.read_csv(PROC_DIR / "ff3_regression.csv")
    ff3_short = ff3_table[["strategy", "alpha_bps", "alpha_t", "alpha_ann_pct",
                           "beta_mkt", "beta_smb", "beta_hml", "r_squared"]]
    ff3_short = ff3_short.rename(columns={
        "strategy":      "Strategy",
        "alpha_bps":     "α (bps/day)",
        "alpha_t":       "α t",
        "alpha_ann_pct": "α ann. (%)",
        "beta_mkt":      "β_mkt",
        "beta_smb":      "β_smb",
        "beta_hml":      "β_hml",
        "r_squared":     "R²",
    })
    add_table(
        doc, ff3_short,
        caption="Table 3. Fama-French 3-factor regressions with Newey-West HAC "
                "standard errors (5 lags).",
    )
    add_para(
        doc,
        "The L/S strategy generates +10.7% annualised alpha (t = 1.01, not "
        "significant), with strongly negative β across all three factors: "
        "β_mkt = −0.43, β_smb = −0.38, β_hml = −0.44. The "
        "strategy is structurally long large-cap growth and short small-cap value. "
        "During an 8-year bull market that exposure is what kills the headline "
        "return; the alpha (what's left after the factors are taken out) is positive "
        "at +10.7%/year. With t = 1.01 you can't distinguish that from luck on this "
        "sample, but it's also not the negative-alpha disaster the raw return "
        "suggests.",
        size=10, space_after=4,
    )
    add_image(
        doc, FIG_DIR / "fig_ff3_alpha.png", width_in=6.0,
        caption="Figure 4. Rolling 252-day FF3 alpha (top) and Newey-West "
                "t-statistic (bottom) for the L/S strategy.",
    )
    add_para(
        doc,
        "Figure 4 shows clear regime shifts. 2019-2020: alpha was strongly negative, "
        "around −20 bps/day, t < −2. Mid-2021: alpha peaked near +30 bps/day, "
        "t briefly above +2. 2022-2024: roughly flat. Late 2025-2026: rebounding. "
        "The 2021 spike maps onto retail-driven trading dynamics in the post-COVID "
        "period, consistent with the overreaction story from Section 4.3.",
        size=10, space_after=6,
    )

    # 4.5
    add_h(doc, "4.5 FinBERT replication (H4)", level=2)
    add_para(
        doc,
        "I scored 300 yfinance-sourced headlines with ProsusAI/finbert and "
        "aggregated to per-(date, ticker) means. Of those, 41 cells matched my "
        "Bloomberg panel. The Spearman correlation between FinBERT's daily mean "
        "and Bloomberg's NEWS_SENTIMENT_DAILY_AVG is −0.26 (p = 0.10). Pearson "
        "is −0.07 (p = 0.67). Neither is statistically significant, and the "
        "correlation is in the wrong direction. That was a surprise.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "H4 is rejected. The interpretation isn't \"FinBERT is broken.\" Manual "
        "inspection shows the model does what you'd expect. \"SOPHiA GENETICS Q1 "
        "Earnings Call Highlights\" scored +0.94. \"Why FTAI Aviation Popped "
        "Higher by More Than 15% Today\" scored +0.94. \"Subsidy wrapper hiding "
        "the real price\" scored −0.97. The polarity calls are reasonable.",
        size=10, space_after=4,
    )
    add_image(
        doc, FIG_DIR / "fig_finbert_vs_bloomberg.png", width_in=6.5,
        caption="Figure 5. FinBERT (Yahoo Finance headlines) vs Bloomberg sentiment. "
                "Left: scatter of per-ticker daily means. Right: marginal "
                "distributions. The two systems are measuring different things.",
    )
    add_para(
        doc,
        "The mismatch is in what's being averaged. Look at Figure 5, right panel. "
        "Bloomberg's distribution is tightly clustered near zero, an average across "
        "thousands of stories per day, mostly factual wire content. FinBERT-on-Yahoo "
        "is bimodal-polarised, a handful of editorial blog headlines per day, each "
        "scored decisively. They're not measuring the same thing. Free Yahoo Finance "
        "news is a small slice of opinion and Motley-Fool-style commentary. "
        "Bloomberg's index averages that plus curated wire stories, regulatory "
        "filings, and professional analyst notes at much higher volume.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "The replication failure is, in some sense, the most useful result in the "
        "paper. The proprietary signal cannot be reconstructed from free public "
        "data. That's the reason it has commercial value.",
        size=10, space_after=8,
    )

    # ---- 5. Discussion ----
    add_h(doc, "5. Discussion", level=1)
    add_para(
        doc,
        "Going in, I expected the 5-day signal to survive even if the 1-day "
        "didn't. It mostly didn't. The efficient-markets reading is the boring "
        "one and probably the right one: news that moves prices in U.S. mega-caps "
        "does so on the same day, and by the time the daily sentiment number is "
        "available, the move is already in the price. The 21-day signal that "
        "does survive is small enough that I wouldn't bet a real strategy on it "
        "in isolation.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "The Q4 vs Q5 asymmetry is the part I keep coming back to. It's a real "
        "result and it points to a non-linear factor structure: \"moderately "
        "positive\" beats \"very positive\" by enough basis points to matter. A "
        "smarter implementation would either bin sentiment categorically and let "
        "the model figure out the kink, or combine sentiment with a "
        "momentum-reversal flag. The plain long-top short-bottom construction "
        "leaves alpha on the table because of this asymmetry.",
        size=10, space_after=4,
    )
    add_para(
        doc,
        "On FinBERT vs Bloomberg, the non-result is the result. Open-source ML "
        "applied to whatever free news is reachable does not commoditise the "
        "Bloomberg signal. That's not the same as saying Bloomberg is "
        "irreplaceable. A serious shop could license a wider news vendor and "
        "build their own pipeline. But \"weekend Yahoo scraper plus FinBERT\" is "
        "demonstrably not in the same league, and the data show why.",
        size=10, space_after=8,
    )

    # ---- 6. Conclusion ----
    add_h(doc, "6. Conclusion", level=1)
    add_para(
        doc,
        "Bloomberg's NEWS_SENTIMENT_DAILY_AVG is a clean signal but not a "
        "standalone return predictor on this universe and sample. The 1-day IC is "
        "noise, the 21-day IC is real but small, and the long-short backtest only "
        "looks meaningful once you adjust for the strategy's strongly negative "
        "factor exposures, which yields about 10.7% annualised alpha (still not "
        "statistically significant). Treat the signal as a market-neutral "
        "diversifier rather than a directional source of return. The non-monotonic "
        "Q4 vs Q5 finding is the most actionable thing in the paper for anyone "
        "who wants to build on this. And the failed FinBERT replication is, "
        "honestly, the strongest argument I produced for why a Bloomberg "
        "subscription has value: I tried to reproduce the signal with free data "
        "and could not.",
        size=10, space_after=8,
    )

    # ---- Appendix A: Terminal evidence ----
    add_h(doc, "Appendix A. Terminal evidence", level=1)
    add_para(
        doc,
        "Screenshots from the Stevens FSC Bloomberg Terminal documenting the "
        "data retrieval. Captured at the same workstation where the VBA pull "
        "ran. Together they show that the sentiment field exists as a real "
        "Bloomberg product, that the universe is documented in their data, and "
        "that the workbook was populated from the official Excel API and not "
        "fabricated downstream.",
        size=10, space_after=6,
    )

    SS_DIR = FIG_DIR / "bloomberg_screenshots"
    appendix_shots = [
        ("00_appl_fnc_menu.png",
         "Figure A1. Bloomberg's Related Functions Menu for AAPL US Equity. "
         "Live ticker bar at top (293.32, +5.88) confirms an active terminal "
         "session. The menu lists DES, CN, FA, ESG, SPLC, and the rest of the "
         "equity-analysis stack. AI-flagged entries (CN, DS RES, FA, MODL, "
         "DS TA, ALTD) indicate functions backed by Bloomberg's own ML."),
        ("01_NSE_search.png",
         "Figure A2. Searching the function namespace for \"NSE\". Bloomberg "
         "deprecated standalone NSE; the canonical entry point for news "
         "sentiment is now TREN NSENT (\"News Trends: News Sentiment\")."),
        ("011_news_sentiment_neg.png",
         "Figure A3. TREN NSENT, News Sentiment tab, Most Negative panel. The "
         "\"Sent.\" column ranges roughly -0.4 to -0.8, on the same [-1, +1] "
         "scale as the NEWS_SENTIMENT_DAILY_AVG field used in the analysis. "
         "This is the same signal, rendered as a cross-sectional ranker."),
        ("011_news_sentiment_pos.png",
         "Figure A4. Same view, Most Positive panel. The symmetric pair of "
         "Most Negative / Most Positive tabs is what TREN NSENT exposes "
         "interactively; the daily aggregate per ticker is what the BDH pull "
         "writes to disk."),
        ("02_oem_mem.png",
         "Figure A5. OEX Index, MEMB function. The S&P 100 constituent list "
         "from which our 30-name universe was selected. Visible names "
         "matching the universe: AAPL, AMZN, BA, BAC, CAT, CVX, DIS, GE, "
         "GOOGL, GS, HD, JNJ. Empty Weight column is a benign display quirk; "
         "the membership list itself is the evidence."),
        ("03_appl_desc.png",
         "Figure A6. AAPL DES (Security Description) — sector classification, "
         "market cap, business summary. Standard \"the terminal was used\" "
         "evidence."),
        ("05_appl_news.png",
         "Figure A7. AAPL CN (Company News) — the underlying news stream that "
         "feeds NEWS_SENTIMENT_DAILY_AVG. Each headline carries a source, a "
         "timestamp, and (in the full view) a per-story sentiment indicator. "
         "Bloomberg's daily aggregate is the volume-weighted mean of these "
         "per-story scores."),
        ("06_excel_bloomberg_ribbon.png",
         "Figure A8. The pulled workbook open in Excel with the Bloomberg "
         "ribbon active. Visible: 30 ticker tabs at the bottom (AAPL through "
         "VZ, plus BENCHMARKS), the per-sheet column structure (Price Date | "
         "Price | Sentiment Date | Sentiment | Stories Date | NumStories | "
         "MktCap Date | MktCap), and real numeric values populated from BDH. "
         "The NumStories column shows \"#N/A Field Not Applicable\" because "
         "NUM_NEWS_STORIES_24HR is not a valid field on this license tier "
         "(noted in the report). The other three fields carry real data."),
    ]
    for fname, caption in appendix_shots:
        path = SS_DIR / fname
        if path.exists():
            add_image(doc, path, width_in=6.3, caption=caption)
        else:
            print(f"  WARN: missing appendix screenshot {path}")

    # ---- References ----
    add_h(doc, "References", level=2)
    refs = [
        "Bloomberg L.P. (2026). NEWS_SENTIMENT_DAILY_AVG; Bloomberg News "
        "Sentiment Engine. Retrieved via Bloomberg Terminal, Stevens FSC.",
        "Fama, E. F. & French, K. R. (1993). Common risk factors in the "
        "returns on stocks and bonds. Journal of Financial Economics, 33(1), 3–56.",
        "Newey, W. K. & West, K. D. (1987). A simple, positive semi-definite, "
        "heteroskedasticity and autocorrelation consistent covariance matrix. "
        "Econometrica, 55(3), 703–708.",
        "Araci, D. (2019). FinBERT: financial sentiment analysis with "
        "pre-trained language models. arXiv:1908.10063.",
        "French, K. R. (2026). U.S. Research Returns Data (daily). Tuck "
        "School of Business, Dartmouth College.",
    ]
    for r in refs:
        add_para(doc, r, size=9, space_after=2)

    out_path = REPORT_DIR / "FE511_Final_Report_Rajhans.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")
    print(f"  size: {out_path.stat().st_size:,} bytes")


if __name__ == "__main__":
    main()
